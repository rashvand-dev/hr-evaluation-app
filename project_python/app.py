import streamlit as st
import pandas as pd
from openai import OpenAI
import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def init_db():
    if 'db' not in st.session_state:
        st.session_state.db = pd.DataFrame({
            # st.session_state.db  == st.session_state['db']
            # 'db' is the key name for our database
            "Name": ["علی رضایی", "سارا محمدی", "رضا کریمی", "مریم احمدی",
                     "سینا طاهری", "فاطمه حسینی", "امید نوری", "زهرا کمالی",
                     "حامد رحیمی", "ندا سلطانی"],
            "Role": ["برنامه‌نویس ارشد Python", "توسعه‌دهنده فرانت‌اند", "مدیر محصول", "برنامه‌نویس جونیور Python",
                     "متخصص داده (Data Scientist)", "طراح رابط کاربری (UI/UX)", "مهندس دوآپس (DevOps)", "مدیر پروژه",
                     "برنامه‌نویس فول‌استک", "تحلیلگر سیستم"],
            "Experience": [5, 2, 4, 1, 3, 4, 2, 5, 3, 4],
            "Quality": [92, 65, 88, 70, 85, 95, 60, 90, 75, 82],
            "Teamwork": [95, 60, 90, 75, 60, 95, 65, 92, 60, 80],
            "Responsibility": [90, 70, 85, 75, 80, 95, 60, 88, 70, 85],
            "Key_Skills": ["پایتون، جنگو، معماری سیستم", "ری‌اکت، جاوا اسکریپت", "مدیریت چابک، اسکرام",
                           "پایتون مقدماتی، گیت",
                           "یادگیری ماشین، SQL", "فیگما، تجربه کاربری", "داکر، کوبرنتیز", "برنامه‌ریزی استراتژیک",
                           "نود.جی‌اس، مونگودبی", "تحلیل نیازمندی‌ها"]
        })


def calculate_weighted_performance(q, t, r, w_q, w_t, w_r):
    total_weight = w_q + w_t + w_r
    if total_weight == 0:
        return 0.0, "Error"

    avg = ((q * w_q) + (t * w_t) + (r * w_r)) / total_weight

    if avg >= 80:
        status = "عالی"
    elif avg >= 60:
        status = "متوسط"
    else:
        status = "نیازمند بهبود"

    return float(avg), status


def generate_ai_report(name, role, exp, skills, average, status):
    prompt = f"""
    مشخصات کارمند:
    - نام: {name}
    - سمت شغلی: {role}
    - سابقه کار: {exp} سال
    - مهارت‌های کلیدی: {skills}
    - نمره عملکرد محاسبه‌شده: {average:.2f} از ۱۰۰
    - وضعیت ارزیابی: {status}

    تو مدیر منابع انسانی (HR) هستی. یک برنامه توسعه فردی (Action Plan) سه‌مرحله‌ای و کاملاً تخصصی متناسب با "سمت شغلی" و "مهارت‌های" این فرد بنویس.
    قوانین اکید:
    ۱. کاملاً بر اساس سمت شغلی ({role}) و تخصص‌هایش صحبت کن.
    ۲. از کلی‌گویی پرهیز کن و راهکارهای فنی و تکنیکال پیشنهاد بده.
    ۳. خروجی را با ساختار تمیز و تیتربندی مناسب (استفاده از علامت # برای تیترها و بولت‌پوینت) به زبان فارسی بنویس.
    """

    # Try to connect to AI. If it fails, don't crash the app.
    try:
        client = OpenAI(
            base_url="https://api.gapgpt.app/v1",
            api_key = st.secrets["MY_API_KEY"]
            # Fix timeout issue on Streamlit Cloud
            timeout=60.0
        )

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system",
                 "content": "تو یک مدیر منابع انسانی در یک شرکت فناوری هستی که گزارش‌های دقیق به زبان فارسی می‌نویسد."},
                {"role": "user", "content": prompt}
            ]
        )

        ai_message = response.choices[0].message.content
        ai_report = str(ai_message) if ai_message is not None else "پاسخی دریافت نشد."

    # If there is an error, save the error message here.
    except Exception as e:
        ai_report = f"خطا در ارتباط با ماژول هوش مصنوعی. خطا: {e}"

    return ai_report


def main():
    st.set_page_config(page_title="سیستم پیشرفته منابع انسانی", page_icon="logo.png", layout="wide")
    init_db()

# UI Fixes: RTL & Mobile Layout
    st.markdown("""
        <style>
        .block-container {
            direction: rtl;
        }
        .ai-output-box {
            direction: rtl;
            text-align: right;
            background-color: #f8f9fa;
            padding: 20px;
            border-radius: 10px;
            border: 1px solid #e0e0e0;
            color: #111111;
            font-family: Tahoma, sans-serif;
            line-height: 1.8;
        }
        .ai-output-box h1, .ai-output-box h2, .ai-output-box h3, .ai-output-box h4 {
            text-align: right;
            color: #0b5ed7;
        }
        .ai-output-box ul, .ai-output-box ol {
            text-align: right;
            padding-right: 20px;
        }
        @media (max-width: 768px) {
            [data-testid="column"] {
                width: 100% !important;
                flex: 1 1 100% !important;
                min-width: 100% !important;
            }
        }
        </style>
    """, unsafe_allow_html=True)

    # تنظیم ستون‌ها برای بزرگ‌تر کردن و وسط‌چین کردن دقیق لوگو
    col_l, col_c, col_r = st.columns([1, 1, 1])
    with col_c:
        st.image("project_python/logo.png", use_container_width=True)

    # وسط‌نویس کردن عنوان اصلی صفحه
    st.markdown("<h1 style='text-align: center;'>سیستم هوشمند ارزیابی عملکرد منابع انسانی</h1>", unsafe_allow_html=True)
    st.markdown("---")

    st.sidebar.header("وزن‌دهی شاخص‌های ارزیابی")
    st.sidebar.caption("تنظیم اهمیت هر متریک:")
    w_quality = st.sidebar.slider("وزن کیفیت کار", 1, 10, 5)
    w_teamwork = st.sidebar.slider("وزن کار تیمی", 1, 10, 3)
    w_resp = st.sidebar.slider("وزن مسئولیت‌پذیری", 1, 10, 2)

    st.sidebar.markdown("---")
    mode = st.sidebar.radio("ناوبری:", ("صفحه وضعیت ارزیابی کارکنان", "مدیریت پایگاه داده", "ورود دستی"))

    if mode == "مدیریت پایگاه داده":
        st.markdown("<h3 style='text-align: center;'>مدیریت پایگاه داده</h3>", unsafe_allow_html=True)
        st.dataframe(st.session_state.db, use_container_width=True)

        st.markdown("**افزودن / ویرایش کارمند:**")
        col1, col2, col3 = st.columns(3)
        new_name = col1.text_input("نام")
        new_role = col2.text_input("سمت شغلی")
        new_exp = col3.number_input("سابقه کار (سال)", 0, 30, 2)

        col4, col5, col6, col7 = st.columns(4)
        new_q = col4.number_input("کیفیت", 0, 100, 50)
        new_t = col5.number_input("کار تیمی", 0, 100, 50)
        new_r = col6.number_input("مسئولیت‌پذیری", 0, 100, 50)
        new_skills = col7.text_input("مهارت‌های کلیدی")

        if st.button("ذخیره در پایگاه داده"):
            if new_name:
                st.session_state.db = st.session_state.db[st.session_state.db["Name"] != new_name]
                new_row = pd.DataFrame({
                    "Name": [new_name], "Role": [new_role], "Experience": [new_exp],
                    "Quality": [new_q], "Teamwork": [new_t], "Responsibility": [new_r], "Key_Skills": [new_skills]
                })
                st.session_state.db = pd.concat([st.session_state.db, new_row], ignore_index=True)
                st.success(f"{new_name} با موفقیت ذخیره شد!")
                st.rerun()
            else:
                st.error("نام نمی‌تواند خالی باشد.")

        st.markdown("**حذف کارمند:**")
        del_name = st.selectbox("انتخاب برای حذف", [""] + st.session_state.db["Name"].tolist())
        if st.button("حذف"):
            if del_name:
                st.session_state.db = st.session_state.db[st.session_state.db["Name"] != del_name]
                st.warning(f"{del_name} حذف شد.")
                st.rerun()

    else:
        found = False
        name, role, exp, skills, quality, teamwork, responsibility = "", "", 0, "", 0, 0, 0

        if mode == "صفحه وضعیت ارزیابی کارکنان":
            st.markdown("<h3 style='text-align: center;'>صفحه وضعیت ارزیابی کارکنان</h3>", unsafe_allow_html=True)
            name_input = st.selectbox("انتخاب کارمند:", [""] + st.session_state.db["Name"].tolist())
            if name_input:
                row = st.session_state.db[st.session_state.db["Name"] == name_input].iloc[0]
                name = str(row["Name"])
                role = str(row["Role"])
                exp = int(row["Experience"])
                skills = str(row["Key_Skills"])
                quality = int(row["Quality"])
                teamwork = int(row["Teamwork"])
                responsibility = int(row["Responsibility"])
                found = True

        elif mode == "ورود دستی":
            st.markdown("<h3 style='text-align: center;'>ارزیابی دستی</h3>", unsafe_allow_html=True)
            name = st.text_input("نام کارمند:")
            role = st.text_input("سمت شغلی:")
            exp = st.number_input("سابقه کار (سال):", 0, 30, 1)
            skills = st.text_input("مهارت‌های کلیدی:")

            col1, col2, col3 = st.columns(3)
            quality = int(col1.number_input("کیفیت", 0, 100, 50))
            teamwork = int(col2.number_input("کار تیمی", 0, 100, 50))
            responsibility = int(col3.number_input("مسئولیت‌پذیری", 0, 100, 50))
            if name:
                found = True

        if not found:
            st.info("لطفاً یک کارمند را انتخاب کنید یا اطلاعات را به صورت دستی وارد نمایید.")
            st.stop()

        st.info(
            f"**سمت شغلی:** {role} &nbsp;&nbsp;|&nbsp;&nbsp; **سابقه کار:** {exp} سال &nbsp;&nbsp;|&nbsp;&nbsp; **مهارت‌ها:** {skills}")

        if st.button("ارزیابی و تولید گزارش هوش مصنوعی", type="primary"):
            team_avg = float(st.session_state.db[["Quality", "Teamwork", "Responsibility"]].mean().mean())

            average, status = calculate_weighted_performance(
                float(quality), float(teamwork), float(responsibility),
                float(w_quality), float(w_t := w_teamwork), float(w_resp)
            )

            st.divider()
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("کیفیت کار", quality)
            c2.metric("کار تیمی", teamwork)
            c3.metric("مسئولیت‌پذیری", responsibility)

            diff = average - team_avg
            c4.metric("میانگین وزن‌دار", f"{average:.2f}", f"{diff:+.1f} نسبت به میانگین تیم")

            if status == "عالی":
                st.success(f"وضعیت: {status}")
            elif status == "متوسط":
                st.warning(f"وضعیت: {status}")
            else:
                st.error(f"وضعیت: {status}")

            with st.spinner('در حال تولید گزارش تخصصی...'):
                ai_report = generate_ai_report(name, role, exp, skills, average, status)
                st.markdown("### تحلیل تخصصی و برنامه توسعه فردی هوش مصنوعی")

                st.markdown(f'<div class="ai-output-box">{ai_report}</div>', unsafe_allow_html=True)

                # تولید فایل Word برای دانلود بهتر زبان فارسی
                doc = Document()

                heading = doc.add_heading('گزارش رسمی ارزیابی عملکرد', level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                # درج اطلاعات کارمند
                info_text = f"نام کارمند: {name}\nسمت: {role}\nسابقه: {exp} سال\nمهارت‌ها: {skills}\nنمره نهایی: {average:.2f}\nوضعیت: {status}\n"
                p_info = doc.add_paragraph(info_text)
                p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                doc.add_heading('تحلیل هوش مصنوعی:', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT

                # درج متن تولید شده توسط هوش مصنوعی
                for line in ai_report.split('\n'):
                    p_ai = doc.add_paragraph(line)
                    p_ai.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                # ذخیره در حافظه برای دانلود
                bio = io.BytesIO()
                doc.save(bio)

                st.download_button(
                    label="دانلود گزارش رسمی (Word)",
                    data=bio.getvalue(),
                    file_name=f"{str(name).replace(' ', '_')}_HR_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )


if __name__ == "__main__":
    main()
